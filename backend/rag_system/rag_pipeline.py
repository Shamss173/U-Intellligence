"""
Core RAG pipeline: ingestion, chunking, embedding, FAISS indexing,
retrieval, re-ranking, and generation. Pure Python, no RAG frameworks.
"""
import logging
import hashlib
import json
import os
import re
import time
from pathlib import Path
import sqlite3
import threading

import faiss
import numpy as np

try:
    # When imported as a package (e.g. from backend).
    from .config import (
        BATCH_SIZE,
        CHUNK_OVERLAP,
        CHUNK_SIZE,
        EMBEDDING_DIM,
        FAISS_INDEX_PATH,
        GCP_PROJECT_ID,
        GCP_LOCATION,
        SERVICE_ACCOUNT_PATH,
        GENERATION_MODEL,
        EMBEDDING_MODEL,
        MIN_CHUNK_LENGTH,
        RERANKER_MODEL,
        SUPPORTED_EXTENSIONS,
        TOP_K_RERANK,
        TOP_K_RETRIEVAL,
    )
except ImportError:
    # When running this file directly (python rag_pipeline.py).
    from config import (  # type: ignore
    BATCH_SIZE,
    CHUNK_OVERLAP,
    CHUNK_SIZE,
    EMBEDDING_DIM,
    FAISS_INDEX_PATH,
    GCP_PROJECT_ID,
    GCP_LOCATION,
    SERVICE_ACCOUNT_PATH,
    GENERATION_MODEL,
    EMBEDDING_MODEL,
    MIN_CHUNK_LENGTH,
    RERANKER_MODEL,
    SUPPORTED_EXTENSIONS,
    TOP_K_RERANK,
    TOP_K_RETRIEVAL,
)

logger = logging.getLogger(__name__)

# Persistent ingestion state (used to skip unchanged files across runs)
# Keep it next to the index files so callers don't depend on CWD.
INGESTION_STATE_PATH = str(Path(FAISS_INDEX_PATH).with_name("ingestion_state.json"))
INGESTION_STATE_VERSION = 1

# Ensure FAISS operations are thread-safe.
_faiss_lock = threading.RLock()

# Persist chunk-level metadata in SQLite instead of pickle.
METADATA_DB_PATH = str(Path(FAISS_INDEX_PATH).with_name("metadata.db"))


def _metadata_db_connect() -> sqlite3.Connection:
    conn = sqlite3.connect(METADATA_DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS chunks (
            chunk_id INTEGER PRIMARY KEY,
            text TEXT,
            source_file TEXT,
            folder_path TEXT,
            department_id TEXT,
            page_number INTEGER,
            file_type TEXT,
            source_path TEXT,
            relative_path TEXT
        )
        """
    )
    return conn


def _reset_chunks_table(conn: sqlite3.Connection) -> None:
    conn.execute("DELETE FROM chunks")


def _insert_chunks_metadata_batch(conn: sqlite3.Connection, metadata_rows: list[dict]) -> None:
    """
    Insert chunk metadata rows.

    metadata_rows dict keys follow the historical pickle format, i.e. they include:
    - 'text' (chunk text)
    - 'chunk_id', 'source_file', 'folder_path', 'file_type', 'source_path', 'relative_path'
    - optional 'page_number'
    - optional 'department_id'
    """
    sql = """
        INSERT OR REPLACE INTO chunks (
            chunk_id,
            text,
            source_file,
            folder_path,
            department_id,
            page_number,
            file_type,
            source_path,
            relative_path
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
    """
    params = []
    for m in metadata_rows:
        # Keep nullability permissive so older metadata dict shapes still work.
        chunk_id = int(m.get("chunk_id"))
        params.append(
            (
                chunk_id,
                m.get("text", ""),
                m.get("source_file", ""),
                m.get("folder_path", ""),
                m.get("department_id"),  # may be absent
                m.get("page_number"),  # may be absent
                m.get("file_type", ""),
                m.get("source_path", ""),
                m.get("relative_path", ""),
            )
        )

    if params:
        conn.executemany(sql, params)


def _load_metadata_list_from_db() -> list[dict]:
    with _metadata_db_connect() as conn:
        rows = conn.execute(
            """
            SELECT chunk_id, text, source_file, folder_path, department_id, page_number, file_type,
                   source_path, relative_path
            FROM chunks
            ORDER BY chunk_id ASC
            """
        ).fetchall()

    metadata_list: list[dict] = []
    for r in rows:
        meta: dict = {
            "chunk_id": int(r["chunk_id"]),
            "text": r["text"] or "",
            "source_file": r["source_file"] or "",
            "folder_path": r["folder_path"] or "",
            "file_type": r["file_type"] or "",
            "source_path": r["source_path"] or "",
            "relative_path": r["relative_path"] or "",
        }
        # Preserve historical pickle dict keys: omit fields when NULL so downstream sees identical shapes.
        if r["department_id"] is not None:
            meta["department_id"] = r["department_id"]
        if r["page_number"] is not None:
            meta["page_number"] = int(r["page_number"])
        metadata_list.append(meta)
    return metadata_list

# ---------------------------------------------------------------------------
# NLTK: ensure punkt is available for sentence tokenization
# ---------------------------------------------------------------------------
def _ensure_nltk_punkt():
    import nltk
    try:
        nltk.data.find("tokenizers/punkt")
        nltk.data.find("tokenizers/punkt_tab/english/")
    except LookupError:
        nltk.download("punkt", quiet=True)
        nltk.download("punkt_tab", quiet=True)


# ---------------------------------------------------------------------------
# Document extraction (PDF, DOCX, TXT, Excel)
# ---------------------------------------------------------------------------
def _normalize_text(text: str) -> str:
    """Remove excessive whitespace, preserve sentence boundaries."""
    if not text or not text.strip():
        return ""
    text = re.sub(r"\s+", " ", text.strip())
    return text


def extract_excel(file_path: str, folder_path: str) -> list[dict]:
    """
    Extract cell contents from Excel (.xlsx/.xls).
    Returns list of dicts: { "text": str, "sheet_name": str, "source_file", "folder_path", "file_type" }.
    """
    results: list[dict] = []
    try:
        relative_path = os.path.relpath(file_path, start=folder_path)
    except Exception:
        relative_path = ""

    ext = Path(file_path).suffix.lower()
    if ext == ".xlsx":
        try:
            from openpyxl import load_workbook
        except Exception as e:
            logger.warning("Missing dependency for .xlsx (%s). Install openpyxl.", e)
            return results
        try:
            wb = load_workbook(filename=file_path, read_only=True, data_only=True)
            for ws in wb.worksheets:
                sheet_name = ws.title or "Sheet"
                rows_text: list[str] = []
                for row in ws.iter_rows(values_only=True):
                    if not row:
                        continue
                    cells = [str(c).strip() for c in row if c is not None and str(c).strip() != ""]
                    if not cells:
                        continue
                    rows_text.append(" | ".join(cells))
                text = _normalize_text(f"[Sheet: {sheet_name}] " + "\n".join(rows_text))
                if text and len(text) >= MIN_CHUNK_LENGTH:
                    results.append({
                        "text": text,
                        "sheet_name": sheet_name,
                        "source_file": os.path.basename(file_path),
                        "folder_path": folder_path,
                        "source_path": os.path.abspath(file_path),
                        "relative_path": relative_path,
                        "file_type": "excel",
                    })
        except Exception as e:
            logger.warning("Failed to parse XLSX %s: %s", file_path, e)
        return results

    if ext == ".xls":
        try:
            import xlrd
        except Exception as e:
            logger.warning("Missing dependency for .xls (%s). Install xlrd.", e)
            return results
        try:
            wb = xlrd.open_workbook(file_path)
            for sheet in wb.sheets():
                sheet_name = sheet.name or "Sheet"
                rows_text: list[str] = []
                for r in range(sheet.nrows):
                    row = sheet.row_values(r)
                    cells = [str(c).strip() for c in row if c is not None and str(c).strip() != ""]
                    if not cells:
                        continue
                    rows_text.append(" | ".join(cells))
                text = _normalize_text(f"[Sheet: {sheet_name}] " + "\n".join(rows_text))
                if text and len(text) >= MIN_CHUNK_LENGTH:
                    results.append({
                        "text": text,
                        "sheet_name": sheet_name,
                        "source_file": os.path.basename(file_path),
                        "folder_path": folder_path,
                        "source_path": os.path.abspath(file_path),
                        "relative_path": relative_path,
                        "file_type": "excel",
                    })
        except Exception as e:
            logger.warning("Failed to parse XLS %s: %s", file_path, e)
        return results

    return results


def extract_pdf(file_path: str, folder_path: str) -> list[dict]:
    """
    Extract text and tables from PDF using pdfplumber.
    Returns list of dicts: { "text": str, "page_number": int, "source_file", "folder_path", "file_type" }.
    """
    import pdfplumber
    results = []
    try:
        relative_path = os.path.relpath(file_path, start=folder_path)
    except Exception:
        relative_path = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                page_texts = []
                # Extract text
                text = page.extract_text()
                if text:
                    page_texts.append(_normalize_text(text))
                # Extract tables
                tables = page.extract_tables()
                for table in tables or []:
                    for row in table:
                        if row and any(cell for cell in row if cell):
                            row_str = " | ".join(f"Col{i+1}: {str(cell or '').strip()}" for i, cell in enumerate(row))
                            page_texts.append(_normalize_text(row_str))
                combined = " ".join(page_texts)
                if combined and len(combined) >= MIN_CHUNK_LENGTH:
                    results.append({
                        "text": combined,
                        "page_number": page_num,
                        "source_file": os.path.basename(file_path),
                        "folder_path": folder_path,
                        "source_path": os.path.abspath(file_path),
                        "relative_path": relative_path,
                        "file_type": "pdf",
                    })
                elif not combined and page_texts:
                    # Page had content but too short after normalize — still attach for chunking across pages if needed
                    pass
    except Exception as e:
        logger.warning("Failed to parse PDF %s: %s", file_path, e)
    return results


def extract_docx(file_path: str, folder_path: str) -> list[dict]:
    """
    Extract paragraphs and tables from DOCX using python-docx.
    Returns list of dicts with "text", "source_file", "folder_path", "file_type". No page_number.
    """
    from docx import Document
    results = []
    try:
        relative_path = os.path.relpath(file_path, start=folder_path)
    except Exception:
        relative_path = ""
    try:
        doc = Document(file_path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(_normalize_text(para.text))
        for table in doc.tables:
            for row in table.rows:
                cells = [str(cell.text or "").strip() for cell in row.cells]
                if any(cells):
                    row_str = " | ".join(f"Column{i+1}: {c}" for i, c in enumerate(cells))
                    parts.append(_normalize_text(row_str))
        full_text = " ".join(parts)
        if full_text and len(full_text) >= MIN_CHUNK_LENGTH:
            results.append({
                "text": full_text,
                "source_file": os.path.basename(file_path),
                "folder_path": folder_path,
                "source_path": os.path.abspath(file_path),
                "relative_path": relative_path,
                "file_type": "docx",
            })
    except Exception as e:
        logger.warning("Failed to parse DOCX %s: %s", file_path, e)
    return results


def extract_txt(file_path: str, folder_path: str) -> list[dict]:
    """Extract plain text from TXT file (UTF-8)."""
    results = []
    try:
        relative_path = os.path.relpath(file_path, start=folder_path)
    except Exception:
        relative_path = ""
    try:
        with open(file_path, encoding="utf-8") as f:
            text = f.read()
        text = _normalize_text(text)
        if text and len(text) >= MIN_CHUNK_LENGTH:
            results.append({
                "text": text,
                "source_file": os.path.basename(file_path),
                "folder_path": folder_path,
                "source_path": os.path.abspath(file_path),
                "relative_path": relative_path,
                "file_type": "txt",
            })
    except Exception as e:
        logger.warning("Failed to read TXT %s: %s", file_path, e)
    return results


def extract_file(file_path: str, folder_path: str) -> list[dict]:
    """Dispatch to appropriate extractor by extension. Returns list of page-level or doc-level blocks."""
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_pdf(file_path, folder_path)
    if ext == ".docx":
        return extract_docx(file_path, folder_path)
    if ext == ".txt":
        return extract_txt(file_path, folder_path)
    if ext in (".xlsx", ".xls"):
        return extract_excel(file_path, folder_path)
    return []


# ---------------------------------------------------------------------------
# Sentence-based chunking (NLTK)
# ---------------------------------------------------------------------------
def chunk_text_into_sentences(full_text: str) -> list[str]:
    """Split text into sentences using NLTK; return list of sentence strings."""
    import nltk
    return nltk.sent_tokenize(full_text)


def build_chunks_from_blocks(blocks: list[dict], starting_chunk_id: int = 0) -> list[dict]:
    """
    From a list of blocks (each with "text" and metadata), build sentence-based chunks
    with CHUNK_SIZE and CHUNK_OVERLAP. Each chunk gets metadata: source_file, folder_path,
    page_number (if present), chunk_id, file_type.
    """
    _ensure_nltk_punkt()
    chunks = []
    chunk_id = int(starting_chunk_id or 0)
    for block in blocks:
        text = block.get("text", "")
        if not text or len(text) < MIN_CHUNK_LENGTH:
            continue
        sentences = chunk_text_into_sentences(text)
        if not sentences:
            continue
        step = max(1, CHUNK_SIZE - CHUNK_OVERLAP)
        for i in range(0, len(sentences), step):
            window = sentences[i : i + CHUNK_SIZE]
            chunk_text = " ".join(window)
            if len(chunk_text) < MIN_CHUNK_LENGTH:
                continue
            meta = {
                "source_file": block.get("source_file", ""),
                "folder_path": block.get("folder_path", ""),
                "chunk_id": chunk_id,
                "file_type": block.get("file_type", ""),
                "source_path": block.get("source_path", ""),
                "relative_path": block.get("relative_path", ""),
            }
            if "page_number" in block:
                meta["page_number"] = block["page_number"]
            chunks.append({"text": chunk_text, "metadata": meta})
            chunk_id += 1
    return chunks


# ---------------------------------------------------------------------------
# Collect files from multiple folders (recursive, deduplicated)
# ---------------------------------------------------------------------------
def collect_files_from_folders(folder_paths: list[str]) -> list[tuple[str, str]]:
    """
    Walk each folder recursively; collect (absolute_path, folder_path) for
    .pdf, .docx, .txt. Deduplicate by absolute path. folder_path is the root
    folder that was passed in (for traceability).
    """
    seen = set()
    collected = []
    for folder in folder_paths:
        folder = os.path.abspath(folder)
        if not os.path.isdir(folder):
            logger.warning("Folder does not exist, skipping: %s", folder)
            continue
        for root, _dirs, files in os.walk(folder):
            for name in files:
                ext = Path(name).suffix.lower()
                if ext not in SUPPORTED_EXTENSIONS:
                    continue
                abs_path = os.path.abspath(os.path.join(root, name))
                if abs_path in seen:
                    continue
                seen.add(abs_path)
                collected.append((abs_path, folder))
    return collected


# ---------------------------------------------------------------------------
# Ingestion state (skip unchanged files across runs)
# ---------------------------------------------------------------------------

def _sha256_file(file_path: str) -> str:
    h = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def _file_fingerprint(file_path: str) -> dict:
    st = os.stat(file_path)
    return {
        "size": int(st.st_size),
        "mtime_ns": int(getattr(st, "st_mtime_ns", int(st.st_mtime * 1e9))),
        "sha256": _sha256_file(file_path),
    }


def _load_ingestion_state() -> dict:
    if not os.path.isfile(INGESTION_STATE_PATH):
        return {"version": INGESTION_STATE_VERSION, "embedding_model": EMBEDDING_MODEL, "index_dim": None, "files": {}}
    try:
        with open(INGESTION_STATE_PATH, "r", encoding="utf-8") as f:
            state = json.load(f)
        if not isinstance(state, dict):
            raise ValueError("bad state type")
        state.setdefault("version", INGESTION_STATE_VERSION)
        state.setdefault("embedding_model", EMBEDDING_MODEL)
        state.setdefault("index_dim", None)
        state.setdefault("files", {})
        if not isinstance(state.get("files"), dict):
            state["files"] = {}
        return state
    except Exception as e:
        logger.warning("Failed to read %s (%s). Rebuilding state from scratch.", INGESTION_STATE_PATH, e)
        return {"version": INGESTION_STATE_VERSION, "embedding_model": EMBEDDING_MODEL, "index_dim": None, "files": {}}


def _save_ingestion_state(state: dict) -> None:
    tmp = INGESTION_STATE_PATH + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(state, f, indent=2, sort_keys=True)
    os.replace(tmp, INGESTION_STATE_PATH)


# ---------------------------------------------------------------------------
# Ingestion: parse files -> blocks -> chunks
# ---------------------------------------------------------------------------
def ingest_files(files_list: list[tuple[str, str]], starting_chunk_id: int = 0) -> list[dict]:
    """
    Ingest supported files from files_list (abs_path, folder_root); return list of chunks
    with "text" and "metadata".
    """
    per_folder: dict[str, int] = {}
    for _abs_path, folder in files_list:
        per_folder[folder] = per_folder.get(folder, 0) + 1
    for folder, count in per_folder.items():
        logger.info("Folder %s: %d file(s) found.", folder, count)

    all_blocks: list[dict] = []
    for abs_path, folder_path in files_list:
        blocks = extract_file(abs_path, folder_path)
        all_blocks.extend(blocks)

    chunks = build_chunks_from_blocks(all_blocks, starting_chunk_id=starting_chunk_id)
    logger.info("Total chunks created: %d", len(chunks))
    return chunks


def ingest_folders(folder_paths: list[str]) -> list[dict]:
    """
    Ingest all supported files from folder_paths; return list of chunks
    with "text" and "metadata" (source_file, folder_path, page_number?, chunk_id, file_type).
    """
    files_list = collect_files_from_folders(folder_paths)
    return ingest_files(files_list, starting_chunk_id=0)


# ---------------------------------------------------------------------------
# Embedding (Vertex AI)
# ---------------------------------------------------------------------------
def _get_vertex_credentials():
    """Load service account credentials for Vertex AI."""
    import google.auth
    import google.oauth2.service_account as sa
    scopes = ["https://www.googleapis.com/auth/cloud-platform"]
    credentials = sa.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_PATH, scopes=scopes
    )
    return credentials


def embed_texts(texts: list[str], task_type: str = "RETRIEVAL_DOCUMENT") -> np.ndarray:
    """
    Embed a list of texts via Vertex AI REST using service account credentials.
    Avoids importing google.genai.types (which crashes on Python 3.14 via PIL).
    """
    import json as _json
    import urllib.request
    from urllib.error import HTTPError
    credentials = _get_vertex_credentials()

    all_embeddings = []
    max_retries = 5
    base_backoff = 5  # seconds
    for i in range(0, len(texts), BATCH_SIZE):
        # Refresh token as needed right before each batch request.
        if not credentials.valid:
            import google.auth.transport.requests
            credentials.refresh(google.auth.transport.requests.Request())

        batch = texts[i: i + BATCH_SIZE]
        instances = [{"content": t, "task_type": task_type} for t in batch]
        body = _json.dumps({"instances": instances}).encode("utf-8")
        url = (
            f"https://{GCP_LOCATION}-aiplatform.googleapis.com/v1/projects/{GCP_PROJECT_ID}"
            f"/locations/{GCP_LOCATION}/publishers/google/models/{EMBEDDING_MODEL}:predict"
        )
        req = urllib.request.Request(url, data=body, method="POST")
        req.add_header("Authorization", f"Bearer {credentials.token}")
        req.add_header("Content-Type", "application/json")

        # Simple retry with backoff for rate limiting
        attempt = 0
        while True:
            try:
                with urllib.request.urlopen(req) as resp:
                    result = _json.loads(resp.read())
                for pred in result.get("predictions", []):
                    emb = pred.get("embeddings", {}).get("values")
                    if emb:
                        all_embeddings.append(emb)
                    else:
                        raise ValueError("No embedding values in prediction")
                break  # batch succeeded
            except HTTPError as e:
                if e.code == 429 and attempt < max_retries:
                    wait_s = base_backoff * (2 ** attempt)
                    logger.warning(
                        "Vertex AI rate limit (429) on batch %d-%d / %d. "
                        "Retrying in %d seconds (attempt %d/%d).",
                        i,
                        min(i + BATCH_SIZE, len(texts)),
                        len(texts),
                        wait_s,
                        attempt + 1,
                        max_retries,
                    )
                    time.sleep(wait_s)
                    attempt += 1
                    continue
                logger.exception("Vertex AI HTTP error for batch %d-%d / %d: %s",
                                 i, min(i + BATCH_SIZE, len(texts)), len(texts), e)
                raise
            except Exception as e:
                logger.exception("Vertex AI embedding failed for batch %d-%d / %d: %s",
                                 i, min(i + BATCH_SIZE, len(texts)), len(texts), e)
                raise
        if i + BATCH_SIZE < len(texts):
            time.sleep(1)
        logger.info("Embedded batch %d-%d / %d", i, min(i + BATCH_SIZE, len(texts)), len(texts))
    return np.array(all_embeddings, dtype=np.float32)


# ---------------------------------------------------------------------------
# FAISS index build and load
# ---------------------------------------------------------------------------
def build_index(folder_paths: list[str]) -> None:
    """
    Ingest folders -> chunk -> embed -> build or update FAISS index (IndexFlatIP with
    L2-normalized vectors). Saves index to FAISS_INDEX_PATH and chunk metadata to metadata.db.

    Robust de-duplication across repeated ingests:
    - Unchanged files are skipped (based on SHA-256 fingerprint).
    - New files are appended to the existing index.
    - If a previously indexed file changed, we rebuild the full index to avoid duplicates/stale vectors.
    """
    files_list = collect_files_from_folders(folder_paths)
    if not files_list:
        logger.warning("No supported files found. Index not built.")
        return

    # Load persistent ingestion state (used to skip unchanged files across runs)
    state = _load_ingestion_state()
    state_files: dict = state.get("files", {}) if isinstance(state.get("files"), dict) else {}

    # Fingerprint current files and detect changes
    current_fingerprints: dict[str, dict] = {}
    unchanged_files: list[tuple[str, str]] = []
    new_files: list[tuple[str, str]] = []
    changed_existing_files: list[tuple[str, str]] = []

    for abs_path, folder_root in files_list:
        try:
            fp = _file_fingerprint(abs_path)
        except Exception as e:
            logger.warning("Failed to fingerprint %s (%s). Skipping file.", abs_path, e)
            continue
        current_fingerprints[os.path.abspath(abs_path)] = fp

        prev = state_files.get(os.path.abspath(abs_path))
        if prev and isinstance(prev, dict) and prev.get("sha256") == fp.get("sha256"):
            unchanged_files.append((abs_path, folder_root))
        else:
            new_files.append((abs_path, folder_root))
            if prev:
                changed_existing_files.append((abs_path, folder_root))

    # Detect deletions: if a previously indexed file is now missing, rebuild.
    removed_prev_files = []
    if state_files:
        for prev_path in list(state_files.keys()):
            if prev_path not in current_fingerprints:
                removed_prev_files.append(prev_path)

    index_exists = os.path.isfile(FAISS_INDEX_PATH) and os.path.isfile(METADATA_DB_PATH)

    # If we have no previous state or no index on disk, do a full rebuild
    if not state_files or not index_exists:
        logger.info("Building full index (%d file(s) scanned).", len(current_fingerprints))
        chunks = ingest_files(files_list, starting_chunk_id=0)
        if not chunks:
            logger.warning("No chunks produced. Index not built.")
            return

        texts = [c["text"] for c in chunks]
        logger.info("Embedding %d chunks in batches of %d...", len(texts), BATCH_SIZE)
        embeddings = embed_texts(texts, task_type="RETRIEVAL_DOCUMENT")
        if len(embeddings) != len(chunks):
            raise RuntimeError("Embedding count mismatch")

        embeddings = np.ascontiguousarray(embeddings.astype(np.float32))
        actual_dim = embeddings.shape[1]
        if actual_dim != EMBEDDING_DIM:
            logger.warning(
                "Embedding dimension from API (%d) differs from config EMBEDDING_DIM (%d). Using %d.",
                actual_dim, EMBEDDING_DIM, actual_dim,
            )
        index = faiss.IndexFlatIP(actual_dim)
        faiss.normalize_L2(embeddings)
        with _faiss_lock:
            index.add(embeddings)

        metadata_rows = []
        for c in chunks:
            m = dict(c["metadata"])
            m["text"] = c["text"]
            metadata_rows.append(m)

        with _faiss_lock:
            faiss.write_index(index, FAISS_INDEX_PATH)
        with _metadata_db_connect() as conn:
            _reset_chunks_table(conn)
            _insert_chunks_metadata_batch(conn, metadata_rows)

        state = {
            "version": INGESTION_STATE_VERSION,
            "embedding_model": EMBEDDING_MODEL,
            "index_dim": int(index.d),
            "files": current_fingerprints,
        }
        _save_ingestion_state(state)

        logger.info(
            "FAISS index saved: %s (%d vectors), metadata: %s",
            FAISS_INDEX_PATH,
            index.ntotal,
            METADATA_DB_PATH,
        )
        return

    # If any previously indexed file was deleted, rebuild to drop stale vectors.
    if removed_prev_files:
        logger.info(
            "Detected %d deleted previously indexed file(s). Rebuilding full index to drop stale vectors.",
            len(removed_prev_files),
        )
        chunks = ingest_files(files_list, starting_chunk_id=0)
        if not chunks:
            logger.warning("No chunks produced. Index not built.")
            return

        texts = [c["text"] for c in chunks]
        logger.info("Embedding %d chunks in batches of %d...", len(texts), BATCH_SIZE)
        embeddings = embed_texts(texts, task_type="RETRIEVAL_DOCUMENT")
        if len(embeddings) != len(chunks):
            raise RuntimeError("Embedding count mismatch")

        embeddings = np.ascontiguousarray(embeddings.astype(np.float32))
        actual_dim = embeddings.shape[1]
        if actual_dim != EMBEDDING_DIM:
            logger.warning(
                "Embedding dimension from API (%d) differs from config EMBEDDING_DIM (%d). Using %d.",
                actual_dim, EMBEDDING_DIM, actual_dim,
            )
        index = faiss.IndexFlatIP(actual_dim)
        faiss.normalize_L2(embeddings)
        with _faiss_lock:
            index.add(embeddings)

        metadata_rows = []
        for c in chunks:
            m = dict(c["metadata"])
            m["text"] = c["text"]
            metadata_rows.append(m)

        with _faiss_lock:
            faiss.write_index(index, FAISS_INDEX_PATH)
        with _metadata_db_connect() as conn:
            _reset_chunks_table(conn)
            _insert_chunks_metadata_batch(conn, metadata_rows)

        state = {
            "version": INGESTION_STATE_VERSION,
            "embedding_model": EMBEDDING_MODEL,
            "index_dim": int(index.d),
            "files": current_fingerprints,
        }
        _save_ingestion_state(state)

        logger.info("FAISS index rebuilt and saved: %s (%d vectors)", FAISS_INDEX_PATH, index.ntotal)
        return

    # Nothing new: skip work entirely
    if not new_files and not changed_existing_files:
        logger.info("No new/changed files detected. Skipping ingest.")
        return

    # If an already-indexed file changed, rebuild to avoid duplicates/stale vectors
    if changed_existing_files:
        logger.info(
            "Detected %d changed previously indexed file(s). Rebuilding full index to avoid duplicates.",
            len(changed_existing_files),
        )
        chunks = ingest_files(files_list, starting_chunk_id=0)
        if not chunks:
            logger.warning("No chunks produced. Index not built.")
            return

        texts = [c["text"] for c in chunks]
        logger.info("Embedding %d chunks in batches of %d...", len(texts), BATCH_SIZE)
        embeddings = embed_texts(texts, task_type="RETRIEVAL_DOCUMENT")
        if len(embeddings) != len(chunks):
            raise RuntimeError("Embedding count mismatch")

        embeddings = np.ascontiguousarray(embeddings.astype(np.float32))
        actual_dim = embeddings.shape[1]
        if actual_dim != EMBEDDING_DIM:
            logger.warning(
                "Embedding dimension from API (%d) differs from config EMBEDDING_DIM (%d). Using %d.",
                actual_dim, EMBEDDING_DIM, actual_dim,
            )
        index = faiss.IndexFlatIP(actual_dim)
        faiss.normalize_L2(embeddings)
        with _faiss_lock:
            index.add(embeddings)

        metadata_rows = []
        for c in chunks:
            m = dict(c["metadata"])
            m["text"] = c["text"]
            metadata_rows.append(m)

        with _faiss_lock:
            faiss.write_index(index, FAISS_INDEX_PATH)
        with _metadata_db_connect() as conn:
            _reset_chunks_table(conn)
            _insert_chunks_metadata_batch(conn, metadata_rows)

        state_files.update(current_fingerprints)
        state = {
            "version": INGESTION_STATE_VERSION,
            "embedding_model": EMBEDDING_MODEL,
            "index_dim": int(index.d),
            "files": state_files,
        }
        _save_ingestion_state(state)

        logger.info("FAISS index rebuilt and saved: %s (%d vectors)", FAISS_INDEX_PATH, index.ntotal)
        return

    # Append only NEW files (unchanged files are skipped)
    logger.info("Appending %d new file(s); skipping %d unchanged file(s).", len(new_files), len(unchanged_files))

    # Load existing index (metadata comes from SQLite)
    with _faiss_lock:
        index = faiss.read_index(FAISS_INDEX_PATH)
    with _metadata_db_connect() as conn:
        max_chunk_id_row = conn.execute("SELECT COALESCE(MAX(chunk_id), -1) AS max_id FROM chunks").fetchone()
        max_chunk_id = int(max_chunk_id_row["max_id"]) if max_chunk_id_row else -1
    next_chunk_id = max_chunk_id + 1

    new_chunks = ingest_files(new_files, starting_chunk_id=next_chunk_id)
    if not new_chunks:
        logger.warning("No chunks produced from new files. Index unchanged.")
        # Still persist fingerprints so we don't retry the same bad file forever
        for p, fp in current_fingerprints.items():
            state_files[p] = fp
        state["files"] = state_files
        _save_ingestion_state(state)
        return

    texts = [c["text"] for c in new_chunks]
    logger.info("Embedding %d new chunk(s) in batches of %d...", len(texts), BATCH_SIZE)
    embeddings = embed_texts(texts, task_type="RETRIEVAL_DOCUMENT")
    if len(embeddings) != len(new_chunks):
        raise RuntimeError("Embedding count mismatch")

    embeddings = np.ascontiguousarray(embeddings.astype(np.float32))
    actual_dim = embeddings.shape[1]
    if int(index.d) != int(actual_dim):
        logger.warning(
            "Embedding dim (%d) does not match existing index dim (%d). Rebuilding full index.",
            actual_dim, int(index.d),
        )
        # Rebuild full index from the scanned folders to ensure consistency
        chunks = ingest_files(files_list, starting_chunk_id=0)
        if not chunks:
            logger.warning("No chunks produced. Index not built.")
            return
        texts = [c["text"] for c in chunks]
        logger.info("Embedding %d chunks in batches of %d...", len(texts), BATCH_SIZE)
        embeddings = embed_texts(texts, task_type="RETRIEVAL_DOCUMENT")
        embeddings = np.ascontiguousarray(embeddings.astype(np.float32))
        actual_dim = embeddings.shape[1]
        index = faiss.IndexFlatIP(actual_dim)
        faiss.normalize_L2(embeddings)
        with _faiss_lock:
            index.add(embeddings)
        metadata_rows = []
        for c in chunks:
            m = dict(c["metadata"])
            m["text"] = c["text"]
            metadata_rows.append(m)
        with _faiss_lock:
            faiss.write_index(index, FAISS_INDEX_PATH)
        with _metadata_db_connect() as conn:
            _reset_chunks_table(conn)
            _insert_chunks_metadata_batch(conn, metadata_rows)
        state_files.update(current_fingerprints)
        state = {
            "version": INGESTION_STATE_VERSION,
            "embedding_model": EMBEDDING_MODEL,
            "index_dim": int(index.d),
            "files": state_files,
        }
        _save_ingestion_state(state)
        logger.info("FAISS index rebuilt and saved: %s (%d vectors)", FAISS_INDEX_PATH, index.ntotal)
        return

    faiss.normalize_L2(embeddings)
    with _faiss_lock:
        index.add(embeddings)

    metadata_rows = []
    for c in new_chunks:
        m = dict(c["metadata"])
        m["text"] = c["text"]
        metadata_rows.append(m)

    with _faiss_lock:
        faiss.write_index(index, FAISS_INDEX_PATH)
    with _metadata_db_connect() as conn:
        _insert_chunks_metadata_batch(conn, metadata_rows)

    for p, fp in current_fingerprints.items():
        state_files[p] = fp
    state = {
        "version": INGESTION_STATE_VERSION,
        "embedding_model": EMBEDDING_MODEL,
        "index_dim": int(index.d),
        "files": state_files,
    }
    _save_ingestion_state(state)

    logger.info(
        "FAISS index updated: %s (+%d vectors), metadata: %s",
        FAISS_INDEX_PATH,
        len(new_chunks),
        METADATA_DB_PATH,
    )


def load_index() -> tuple[faiss.Index, list[dict]]:
    """
    Load FAISS index and metadata from disk. Raises FileNotFoundError if index missing.
    """
    if not os.path.isfile(FAISS_INDEX_PATH):
        raise FileNotFoundError(f"Index not found at {FAISS_INDEX_PATH}. Please run --ingest first.")
    with _faiss_lock:
        index = faiss.read_index(FAISS_INDEX_PATH)
    if not os.path.isfile(METADATA_DB_PATH):
        raise FileNotFoundError(
            f"Metadata DB not found at {METADATA_DB_PATH}. Please run --ingest first."
        )
    metadata_list = _load_metadata_list_from_db()
    logger.info("Loaded index with %d vectors", index.ntotal)
    return index, metadata_list


# ---------------------------------------------------------------------------
# Retrieval: embed query -> FAISS search -> top-K candidates
# ---------------------------------------------------------------------------
def retrieve(index: faiss.Index, metadata_list: list[dict], query: str, top_k: int = TOP_K_RETRIEVAL) -> list[dict]:
    """
    Embed query, search FAISS, return list of dicts with "text", "metadata", "score".
    """
    query_emb = embed_texts([query], task_type="RETRIEVAL_QUERY")
    query_emb = np.ascontiguousarray(query_emb.astype(np.float32))
    faiss.normalize_L2(query_emb)

    with _faiss_lock:
        scores, indices = index.search(query_emb, min(top_k, index.ntotal))
    candidates = []
    for j, idx in enumerate(indices[0]):
        if idx < 0:
            continue
        meta = metadata_list[idx]
        text = meta.get("text", "")
        score = float(scores[0][j])
        candidates.append({"text": text, "metadata": {k: v for k, v in meta.items() if k != "text"}, "score": score})
    logger.info("Retrieved %d candidates (scores: %s)", len(candidates), [round(c["score"], 4) for c in candidates[:5]])
    return candidates


# ---------------------------------------------------------------------------
# Re-ranking (pure-Python TF-IDF — avoids scikit-learn crash on Python 3.14)
# ---------------------------------------------------------------------------

def _tokenize(text: str) -> list[str]:
    """Simple whitespace + punctuation tokenizer, lowercased."""
    import re as _re
    return _re.findall(r"[a-z0-9]+", text.lower())


def rerank(query: str, candidates: list[dict], top_k: int = TOP_K_RERANK) -> list[dict]:
    """
    Re-rank candidates using TF-IDF cosine similarity (pure Python, no scipy/sklearn).
    Returns top_k chunks sorted by relevance to query.
    """
    import math
    if not candidates:
        return []

    query_tokens = _tokenize(query)
    if not query_tokens:
        return candidates[:top_k]

    # Build corpus token lists
    corpus = [_tokenize(c["text"]) for c in candidates]
    N = len(corpus)

    # Compute IDF for each query term
    idf = {}
    for term in set(query_tokens):
        df = sum(1 for doc in corpus if term in doc)
        idf[term] = math.log((N + 1) / (df + 1)) + 1.0

    def score(doc_tokens: list[str]) -> float:
        tf_map: dict[str, float] = {}
        for t in doc_tokens:
            tf_map[t] = tf_map.get(t, 0) + 1
        total = len(doc_tokens) or 1
        s = 0.0
        for term in query_tokens:
            tf = tf_map.get(term, 0) / total
            s += tf * idf.get(term, 1.0)
        return s

    scored = [(score(doc_tokens), c) for doc_tokens, c in zip(corpus, candidates)]
    scored.sort(key=lambda x: x[0], reverse=True)
    top = [c for _, c in scored[:top_k]]
    logger.info("Re-ranked top %d (TF-IDF scores: %s)", len(top), [round(s, 4) for s, _ in scored[:top_k]])
    return top



# ---------------------------------------------------------------------------
# Generation (Vertex AI)
# ---------------------------------------------------------------------------
def generate_answer(
    query: str,
    context_chunks: list[dict],
    conversation_history: list[dict] = None,
) -> tuple[str, list[dict]]:
    """
    Build context from chunks, call Gemini via Vertex AI REST API,
    return (answer_text, list of source citations).
    Avoids importing google.genai.types (which crashes on Python 3.14 via PIL).
    """
    import json as _json
    import urllib.request

    credentials = _get_vertex_credentials()
    if not credentials.valid:
        import google.auth.transport.requests
        credentials.refresh(google.auth.transport.requests.Request())

    # Build context with explicit source labels so the model can ground and cite properly
    context_parts = []
    for c in context_chunks:
        meta = c.get("metadata", {})
        src = meta.get("source_file", "")
        page = meta.get("page_number")
        page_str = f", page {page}" if page is not None else ""
        context_parts.append(f"[Source: {src}{page_str}]\n{c['text']}")
    context = "\n\n".join(context_parts)

    system_prompt = f"""You are a precise document Q&A assistant. You answer questions using ONLY the provided context from the user's ingested documents (handbooks, policies, HR, finance, etc.).

Instructions:
- The context is made of excerpts from the document(s), often in document order. Multiple excerpts from the same source/page may form one continuous section (e.g. a list "a. ... b. ... c. ..."). Use all excerpts together to answer; if one excerpt ends mid-list, the next excerpt may continue it.
- Answer using exact terms, names, numbers, and details from the context. If the answer is present (including rephrased or synonyms), provide it. Only say "I don't have enough information in the provided documents to answer this." when the context truly does not contain the answer.
- When asked to list, enumerate, or "enlist" items (grades, categories, requirements, steps, points), gather every such item from the entire context and list them completely as in the document.
- Do not invent information. Do not hallucinate. If a list or section is incomplete in the context, say what is present and that the full list may continue elsewhere in the document.

Context:
{context}

Question: {query}

Answer:"""

    citations = []
    for c in context_chunks:
        meta = c.get("metadata", {})
        citations.append({
            "source_file": meta.get("source_file", ""),
            "folder_path": meta.get("folder_path", ""),
            "page_number": meta.get("page_number"),
        })

    # Prepend last 6 conversation messages (if any) before the final user message.
    # Gemini uses `model` for assistant turns.
    contents = []
    if conversation_history:
        last_messages = conversation_history[-6:]
        for msg in last_messages:
            role = msg.get("role")
            gemini_role = "user" if role == "user" else "model"
            contents.append(
                {"role": gemini_role, "parts": [{"text": msg.get("content", "")}]}
            )

    contents.append({"role": "user", "parts": [{"text": system_prompt}]})

    body = _json.dumps({"contents": contents}).encode("utf-8")
    url = (
        f"https://{GCP_LOCATION}-aiplatform.googleapis.com/v1/projects/{GCP_PROJECT_ID}"
        f"/locations/{GCP_LOCATION}/publishers/google/models/{GENERATION_MODEL}:generateContent"
    )
    req = urllib.request.Request(url, data=body, method="POST")
    req.add_header("Authorization", f"Bearer {credentials.token}")
    req.add_header("Content-Type", "application/json")

    try:
        with urllib.request.urlopen(req) as resp:
            result = _json.loads(resp.read())
        answer = result["candidates"][0]["content"]["parts"][0]["text"]
    except Exception as e:
        logger.exception("Vertex AI generation failed: %s", e)
        answer = "An error occurred while generating the answer. Please try again later."
    return answer.strip(), citations


# ---------------------------------------------------------------------------
# Order chunks by document position (so lists and enumerations stay coherent)
# ---------------------------------------------------------------------------
def _sort_chunks_by_document_order(chunks: list[dict]) -> list[dict]:
    """
    Sort chunks by source_file, then page_number, then chunk_id.
    Ensures the model sees context in document order so lists (a., b., c.) aren't split.
    """
    def sort_key(c: dict) -> tuple:
        meta = c.get("metadata", {})
        source = meta.get("source_file", "")
        page = meta.get("page_number")
        page = page if page is not None else 0
        chunk_id = meta.get("chunk_id", 0)
        return (source, page, chunk_id)
    return sorted(chunks, key=sort_key)


# ---------------------------------------------------------------------------
# Full RAG query pipeline
# ---------------------------------------------------------------------------
def query_rag(question: str, conversation_history: list[dict] = None) -> tuple[str, list[dict]]:
    """
    Load index -> embed query -> retrieve top TOP_K_RETRIEVAL -> rerank to TOP_K_RERANK
    -> sort by document order -> generate with Gemini. Returns (answer, citations).
    """
    index, metadata_list = load_index()
    candidates = retrieve(index, metadata_list, question, top_k=TOP_K_RETRIEVAL)
    if not candidates:
        return "No context found to answer the query.", []
    if candidates:
        # BM25 hybrid scoring over the retrieved FAISS candidate set only.
        from rank_bm25 import BM25Okapi

        def _normalize_0_1(values: list[float]) -> list[float]:
            if not values:
                return []
            vmin = float(min(values))
            vmax = float(max(values))
            if vmax == vmin:
                return [0.0 for _ in values]
            return [(float(v) - vmin) / (vmax - vmin) for v in values]

        faiss_scores = [float(c.get("score", 0.0)) for c in candidates]
        bm25_corpus_tokens = [_tokenize(c.get("text", "")) for c in candidates]
        bm25 = BM25Okapi(bm25_corpus_tokens)
        bm25_scores = bm25.get_scores(_tokenize(question))

        faiss_norm = _normalize_0_1(faiss_scores)
        bm25_norm = _normalize_0_1([float(s) for s in bm25_scores])

        for i, c in enumerate(candidates):
            final_score = 0.6 * faiss_norm[i] + 0.4 * bm25_norm[i]
            c["final_score"] = final_score

        logger.info(
            "Hybrid scoring top candidates (faiss_norm, bm25_norm, final): %s",
            [
                (
                    round(faiss_norm[i], 4),
                    round(bm25_norm[i], 4),
                    round(0.6 * faiss_norm[i] + 0.4 * bm25_norm[i], 4),
                )
                for i in range(min(5, len(candidates)))
            ],
        )

        # Re-sort candidates so downstream selection/order reflects hybrid scoring.
        candidates.sort(key=lambda x: x.get("final_score", 0.0), reverse=True)

    top_chunks = rerank(question, candidates, top_k=TOP_K_RERANK)
    top_chunks = _sort_chunks_by_document_order(top_chunks)
    answer, citations = generate_answer(question, top_chunks, conversation_history=conversation_history)
    return answer, citations
